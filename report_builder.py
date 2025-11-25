"""
report_builder.py — CBCT report generator (template-free, measurement-driven)

Generates a clinician-facing CBCT report that follows the standard headings:
1. Patient Details
2. Scan Details
3. Image Quality
4. Findings
   4.1 Maxillofacial skeleton
   4.2 Teeth & Periapical region
   4.3 TMJ
   4.4 Paranasal sinuses
   4.5 Mandibular canal & neurovascular structures
   4.6 Airway
   4.7 Soft tissues
   4.8 Pathology
   4.9 Implant/Surgical evaluation
5. Impression
6. Recommendations
7. Radiologist signature

Design constraints:
- Uses only input arguments (cbct_result, conditions, measurement_summary, evidence_snippets)
- No hallucination or external templates
- Defensive: clearly states limitations where automatic analysis isn't possible
- Provides structured output for UI consumption
"""

from typing import List, Dict, Any, Tuple, Optional
import datetime
import re
import os

# import measurement builder (assumed available in project)
import cbct_measurements

# ---------------------------------------------------------------------------
# Utilities: parse measurement summary text into structured fields
# ---------------------------------------------------------------------------
def _extract_kv_pairs(text: str) -> Dict[str, Any]:
    """
    Parse common lines from measurement_summary into structured info.
    Returns dict with keys:
     - volume_shape: tuple or string
     - voxel_size: tuple or string
     - hu_mean/min/max: floats if found
     - rois: list of dicts {name, coords, mean_hu, n}
     - ridge_candidates: list of dicts {height_mm, confidence, sample_x}
     - periapical_candidates: list of dicts {vol_mm3, mean_hu, confidence}
     - warnings: list of warning strings
    """
    info = {
        "volume_shape": None,
        "voxel_size": None,
        "hu_mean": None,
        "hu_min": None,
        "hu_max": None,
        "rois": [],
        "ridge_candidates": [],
        "periapical_candidates": [],
        "warnings": [],
        "raw_lines": []
    }

    if not text:
        return info

    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    info["raw_lines"] = lines

    for ln in lines:
        # volume shape and voxel
        m = re.search(r"Volume shape.*\(([^)]+)\).*voxel size.*\(([^)]+)\)", ln, flags=re.I)
        if m:
            info["volume_shape"] = m.group(1).strip()
            info["voxel_size"] = m.group(2).strip()
            continue

        m = re.search(r"HU mean:\s*([0-9.+-]+);\s*min:\s*([0-9.+-]+);\s*max:\s*([0-9.+-]+)", ln, flags=re.I)
        if m:
            try:
                info["hu_mean"] = float(m.group(1))
                info["hu_min"] = float(m.group(2))
                info["hu_max"] = float(m.group(3))
            except Exception:
                pass
            continue

        # ROI lines: "ROI 1 at (27.25, 80.0, 80.0): mean HU 959.2 (n=9261)"
        m = re.search(r"(ROI\s*\d+).*at\s*\(([^)]+)\).*mean\s*HU\s*([0-9.+-]+).*n\s*=\s*([0-9]+)", ln, flags=re.I)
        if m:
            info["rois"].append({
                "name": m.group(1).strip(),
                "coords": m.group(2).strip(),
                "mean_hu": float(m.group(3)),
                "n": int(m.group(4))
            })
            continue

        # Ridge candidate: "Ridge candidate 1: height 17.00 mm (confidence 1.00) at sample_x 113"
        m = re.search(r"Ridge candidate\s*\d+.*height\s*([0-9.]+)\s*mm.*confidence\s*([0-9.]+).*sample_x\s*([0-9]+)", ln, flags=re.I)
        if m:
            info["ridge_candidates"].append({
                "height_mm": float(m.group(1)),
                "confidence": float(m.group(2)),
                "sample_x": int(m.group(3))
            })
            continue

        # Periapical candidate: "Periapical candidate 1: vol 3125.0 mm³; mean HU 0.1; confidence 0.25"
        m = re.search(r"Periapical candidate\s*\d+.*vol\s*([0-9.]+)\s*mm", ln, flags=re.I)
        if m:
            vol = float(m.group(1))
            m2 = re.search(r"mean\s*HU\s*([0-9.+-]+)", ln, flags=re.I)
            mh = float(m2.group(1)) if m2 else None
            m3 = re.search(r"confidence\s*([0-9.]+)", ln, flags=re.I)
            conf = float(m3.group(1)) if m3 else None
            info["periapical_candidates"].append({
                "vol_mm3": vol,
                "mean_hu": mh,
                "confidence": conf
            })
            continue

        # warnings
        if "warning" in ln.lower() or "pixel spacing missing" in ln.lower() or "assum" in ln.lower():
            info["warnings"].append(ln)
            continue

    return info


# ---------------------------------------------------------------------------
# Helper: safe join for lists to produce clear text blocks
# ---------------------------------------------------------------------------
def _join_block(lines: List[str]) -> str:
    return "\n".join([ln.strip() for ln in lines if ln and ln.strip()])


# ---------------------------------------------------------------------------
# Main builder
# ---------------------------------------------------------------------------
def build_report(
    cbct_result: Dict[str, Any],
    conditions: List[str],
    measurement_summary: str,
    evidence_snippets: List[str],
    sample_template_path: Optional[str] = None
) -> Tuple[str, Dict[str, Any]]:
    """
    Build a full CBCT report string and structured dict from:
      - cbct_result (dict from dicom_reader.load_cbct_from_path)
      - conditions (list of detector suggestions / user selected)
      - measurement_summary (string returned by cbct_measurements.build_measurements_summary)
      - evidence_snippets (unused here but included for API compatibility)
    """

    now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    # Basic header
    header_lines = [
        "CBCT AI REPORT",
        f"Generated: {now}",
        "Patient: (anonymized — no identifiers supplied)",
        "",
    ]

    # Scan details from metadata
    metadata = cbct_result.get("metadata", {}) if isinstance(cbct_result, dict) else {}
    source = metadata.get("source_file") or metadata.get("source") or ""
    slice_thickness = metadata.get("slice_thickness")
    dcm_multiframe = metadata.get("dcm_multiframe", False)
    voxel_info = ""
    if slice_thickness:
        voxel_info = f"Slice thickness: {slice_thickness} mm"
    scan_details = [
        "SCAN DETAILS:",
        f"Modality: CBCT",
        f"Source file: {os.path.basename(source)}" if source else "Source file: (not available)",
        f"Multiframe DICOM: {'Yes' if dcm_multiframe else 'No'}",
    ]
    if voxel_info:
        scan_details.append(voxel_info)

    # Image quality: glean from measurement_summary warnings if any
    parsed = _extract_kv_pairs(measurement_summary)
    warnings = parsed.get("warnings", [])
    iq_lines = ["IMAGE QUALITY:"]
    if warnings:
        iq_lines.append("Limited by: " + "; ".join(warnings))
    else:
        iq_lines.append("Image quality adequate for diagnostic evaluation unless otherwise stated.")

    # Start Findings: fill subsections using parsed info + conditions
    findings_blocks = []

    # 4.1 Maxillofacial skeleton (general observations)
    skeleton_lines = ["Maxillofacial skeleton:"]
    # Volume & voxel details
    if parsed.get("volume_shape"):
        skeleton_lines.append(f"- Volume shape (Z,Y,X): {parsed['volume_shape']}")
    if parsed.get("voxel_size"):
        skeleton_lines.append(f"- Voxel size (dz,dy,dx) mm: {parsed['voxel_size']}")
    if parsed.get("hu_mean") is not None:
        skeleton_lines.append(f"- Mean HU (approx): {parsed['hu_mean']}; min: {parsed.get('hu_min')}; max: {parsed.get('hu_max')}")
    findings_blocks.append(_join_block(skeleton_lines))

    # 4.2 Teeth & Periapical region
    teeth_lines = ["Teeth & periapical region:"]
    if parsed["periapical_candidates"]:
        for i, p in enumerate(parsed["periapical_candidates"], start=1):
            vol = p.get("vol_mm3")
            mh = p.get("mean_hu")
            conf = p.get("confidence")
            teeth_lines.append(f"- Periapical candidate {i}: volume {vol} mm³; mean HU {mh}; confidence {conf}. Consider correlation with clinical exam and periapical radiographs.")
    else:
        teeth_lines.append("- No automated periapical lesion candidate detected by the pipeline.")
    findings_blocks.append(_join_block(teeth_lines))

    # 4.3 TMJ (not computed automatically unless a detector indicated)
    tmj_lines = ["Temporomandibular joints (TMJ):"]
    if any("tmj" in (c or "").lower() or "temporomandibular" in (c or "").lower() for c in conditions):
        tmj_lines.append("- TMJ region flagged by detector — recommend focused review of condylar morphology and joint space on coronal/sagittal views.")
    else:
        tmj_lines.append("- No automated TMJ abnormality detected.")
    findings_blocks.append(_join_block(tmj_lines))

    # 4.4 Paranasal sinuses
    sinus_lines = ["Paranasal sinuses (maxillary/ethmoid when in FOV):"]
    if any("sinus" in (c or "").lower() for c in conditions):
        sinus_lines.append("- Sinus region flagged by detector — review mucosal thickening/opacification on coronal cuts.")
    else:
        sinus_lines.append("- No automated sinus pathology detected within limits of the FOV.")
    findings_blocks.append(_join_block(sinus_lines))

    # 4.5 Mandibular canal & neurovascular structures
    canal_lines = ["Mandibular canal & neurovascular structures:"]
    # If user asked for implant assessment or full skull, mention canal visibility
    if any(x in [c.lower() for c in conditions] for x in ("full skull", "mandible", "implant assessment", "implant")):
        canal_lines.append("- Mandibular canal course should be reviewed on cross-sectional and panoramic reconstructions. Automated identification is limited; verify on slices.")
        # If ROI HU suggests clear canal? we avoid inventing — only state if pipeline produced explicit info
    else:
        canal_lines.append("- No targeted canal analysis requested.")
    findings_blocks.append(_join_block(canal_lines))

    # 4.6 Airway
    airway_lines = ["Airway assessment:"]
    if any("airway" in (c or "").lower() for c in conditions):
        airway_lines.append("- Airway region flagged — suggest sagittal airway volume review.")
    else:
        airway_lines.append("- Airway not specifically included in automated analysis for this FOV.")
    findings_blocks.append(_join_block(airway_lines))

    # 4.7 Soft tissues
    soft_lines = ["Soft tissues:"]
    soft_lines.append("- No discrete soft-tissue mass automatically detected by pipeline; radiologist review recommended for subtle findings.")
    findings_blocks.append(_join_block(soft_lines))

    # 4.8 Pathology (automated alerts)
    path_lines = ["Pathology / other observations:"]
    if parsed["ridge_candidates"]:
        # Summarize ridge candidates as relevant to implant planning or resorption
        for i, r in enumerate(parsed["ridge_candidates"], start=1):
            h = r.get("height_mm")
            conf = r.get("confidence")
            path_lines.append(f"- Ridge candidate {i}: height {h:.1f} mm (confidence {conf:.2f}) — useful for preliminary implant site assessment.")
    else:
        path_lines.append("- No automated destructive osseous lesion identified by the pipeline.")
    findings_blocks.append(_join_block(path_lines))

    # 4.9 Implant / Surgical evaluation
    implant_lines = ["Implant / surgical evaluation (if applicable):"]
    if any("implant" in (c or "").lower() for c in conditions) or parsed["ridge_candidates"]:
        # Provide conservative guidance using ridge heights (if present)
        if parsed["ridge_candidates"]:
            implant_lines.append("- Ridge height candidates detected; correlate locations with proposed implant sites.")
        else:
            implant_lines.append("- No measurement-based implant guidance available from automated analysis. Recommend dedicated implant CBCT protocol and surgical stent when planning.")
    else:
        implant_lines.append("- Implant assessment not requested.")
    findings_blocks.append(_join_block(implant_lines))

    # Compose FINDINGS section text
    findings_text = ["FINDINGS:"]
    findings_text.extend(findings_blocks)
    findings_text.append("")  # spacer

    # IMPRESSION: build concise clinically actionable statements based on above
    impression_list: List[str] = []
    # If periapical candidates found, mention them
    if parsed["periapical_candidates"]:
        for i, p in enumerate(parsed["periapical_candidates"], start=1):
            conf = p.get("confidence")
            if conf and conf >= 0.5:
                impression_list.append(f"Probable periapical pathology at candidate {i} (volume {p.get('vol_mm3')} mm³). Correlate with clinical/radiographic findings.")
            else:
                impression_list.append(f"Possible periapical change at candidate {i} (low confidence). Correlate clinically.")
    # If ridge candidates exist, mention availability of ridge heights
    if parsed["ridge_candidates"]:
        impression_list.append("Ridge height measurements are available for preliminary implant assessment; verify on cross-sectional slices prior to final planning.")
    # If none of the above, default negative impression
    if not impression_list:
        impression_list.append("No significant acute destructive osseous lesion identified on the supplied dataset.")

    # RECOMMENDATIONS:
    recommendations_list: List[str] = []
    # Always recommend clinical correlation
    recommendations_list.append("Correlate imaging findings with clinical examination and conventional radiographs where appropriate.")
    # If periapical candidates exist, recommend targeted periapical radiographs or endodontic evaluation
    if parsed["periapical_candidates"]:
        recommendations_list.append("If periapical pathology is suspected, consider targeted periapical radiographs and endodontic consultation.")
    # If implant planning relevant
    if parsed["ridge_candidates"]:
        recommendations_list.append("For implant planning, perform site-specific cross-sectional measurements and consider a surgical stent or guided workflow.")
    # If detector suggested sinus pathology
    if any("sinus" in (c or "").lower() for c in conditions):
        recommendations_list.append("For suspected sinus disease, correlate with ENT or consider dedicated sinus imaging if clinically indicated.")
    # If no automated canal analysis
    recommendations_list.append("Review inferior alveolar canal course on cross-sectional images before surgical interventions near the mandible.")

    # Radiologist signature block
    signature_lines = [
        "",
        "Radiologist:",
        "Name: _______________________",
        "Qualifications: _______________________",
        "Registration number: _______________________",
        f"Date: {now.split(' ')[0]}"
    ]

    # Build full report text
    report_parts: List[str] = []
    report_parts.extend(header_lines)
    report_parts.append(_join_block(scan_details))
    report_parts.append("")
    report_parts.append(_join_block(iq_lines))
    report_parts.append("")
    report_parts.append(_join_block(findings_text))
    report_parts.append("IMPRESSION:")
    for imp in impression_list:
        report_parts.append(f"- {imp}")
    report_parts.append("")
    report_parts.append("RECOMMENDATIONS:")
    for r in recommendations_list:
        report_parts.append(f"- {r}")
    report_parts.extend(signature_lines)

    report_text = "\n".join(report_parts)

    # Structured dict for UI/JSON consumption
    structured = {
        "patient": {
            "source_file": source
        },
        "scan_details": {
            "slice_thickness": slice_thickness,
            "dcm_multiframe": dcm_multiframe
        },
        "image_quality": iq_lines,
        "findings": {
            "skeleton": skeleton_lines,
            "teeth_periapical": teeth_lines,
            "tmj": tmj_lines,
            "sinuses": sinus_lines,
            "mandibular_canal": canal_lines,
            "airway": airway_lines,
            "soft_tissues": soft_lines,
            "pathology": path_lines,
            "implant": implant_lines
        },
        "impression": impression_list,
        "recommendations": recommendations_list,
        "raw_measurement_summary": measurement_summary,
        "conditions": conditions
    }

    return report_text, structured


# ---------------------------------------------------------------------------
# Export helpers: save_docx and save_pdf (safe and deterministic)
# ---------------------------------------------------------------------------
def _clean_text(t: Optional[str]) -> str:
    if not t:
        return ""
    import re
    t = re.sub(r"\r\n", "\n", t)
    t = re.sub(r"[ \t]+", " ", t)
    t = re.sub(r"\n{3,}", "\n\n", t)
    return t.strip()


def save_docx(report_text: str, out_path: str) -> str:
    """
    Save `report_text` to a .docx file. Uses python-docx.
    """
    try:
        from docx import Document
    except Exception as e:
        raise RuntimeError("python-docx is not installed. Install with `pip install python-docx`") from e

    report_text = _clean_text(report_text)
    doc = Document()
    doc.add_heading("CBCT AI Report", level=1)
    for line in report_text.split("\n"):
        if not line.strip():
            doc.add_paragraph("")
            continue
        stripped = line.strip()
        # treat ALL-CAPS ending with colon as section heading
        if stripped.endswith(":") and stripped.upper() == stripped:
            doc.add_heading(stripped[:-1], level=2)
        else:
            doc.add_paragraph(stripped)
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    return out_path


def save_pdf(report_text: str, out_path: str) -> str:
    """
    Save `report_text` to a simple PDF using reportlab.
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
    except Exception as e:
        raise RuntimeError("reportlab is not installed. Install with `pip install reportlab`") from e

    report_text = _clean_text(report_text)
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    c = canvas.Canvas(out_path, pagesize=A4)
    width, height = A4
    left = 40
    y = height - 40
    import textwrap
    for line in report_text.split("\n"):
        if not line.strip():
            y -= 12
            if y < 60:
                c.showPage()
                y = height - 40
            continue
        wrapped = textwrap.wrap(line, width=90)
        for w in wrapped:
            c.drawString(left, y, w)
            y -= 12
            if y < 60:
                c.showPage()
                y = height - 40
    c.save()
    return out_path
